# src/engine/parser.py

import re
from pathlib import Path
import pdfplumber
import docx

KNOWN_SECTIONS = {
    "experience": ["work experience", "professional experience", "experience", "employment history"],
    "skills": ["technical skills", "skills", "core competencies", "technologies"],
    "education": ["education", "academic background"],
    "projects": ["projects", "personal projects", "key projects"]
}

def extract_raw_text_pdf(pdf_path: str) -> str:
    """Extracts text from a PDF file preserving layout order using pdfplumber."""
    text_blocks = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            # Extract text preserving visual spatial layout
            page_text = page.extract_text(layout=True)
            if page_text:
                text_blocks.append(page_text)
    return "\n".join(text_blocks)

def extract_raw_text_docx(docx_path: str) -> str:
    """Extracts text from a .docx file line by line."""
    doc = docx.Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    return "\n".join(full_text)

def classify_section_header(line: str) -> str | None:
    """Detects if a given line is a known resume section header."""
    clean_line = line.strip().lower()
    clean_line = re.sub(r'[^a-z\s]', '', clean_line)
    
    for section_key, aliases in KNOWN_SECTIONS.items():
        if clean_line in aliases:
            return section_key
    return None

def extract_bullet_points(section_text: str) -> list[str]:
    """Splits raw section text into individual clean bullet points."""
    lines = section_text.split("\n")
    bullets = []
    
    for line in lines:
        cleaned = line.strip()
        if not cleaned:
            continue
            
        # Strip common bullet point symbols (•, -, *, numbers)
        cleaned = re.sub(r'^[•\-\*\d\.\)\>]\s*', '', cleaned).strip()
        
        # Only retain valid bullet points (ignoring short lines like dates or job titles)
        if len(cleaned) > 20:
            bullets.append(cleaned)
            
    return bullets

def parse_resume_to_structure(file_path: str) -> dict:
    """
    Main Ingestor function. Takes a PDF or DOCX file path and converts it into 
    a structured Master Candidate Profile JSON.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    # 1. Extract raw text based on extension
    if path.suffix.lower() == ".pdf":
        raw_text = extract_raw_text_pdf(file_path)
    elif path.suffix.lower() in [".docx", ".doc"]:
        raw_text = extract_raw_text_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Please upload .pdf or .docx")

    # 2. Segment text into structured sections
    lines = raw_text.split("\n")
    current_section = "general"
    sections = {
        "general": [],
        "experience": [],
        "skills": [],
        "education": [],
        "projects": []
    }

    for line in lines:
        detected_header = classify_section_header(line)
        if detected_header:
            current_section = detected_header
        else:
            if line.strip():
                sections[current_section].append(line.strip())

    # 3. Process experience & project bullets into structured objects
    exp_bullets = extract_bullet_points("\n".join(sections["experience"]))
    proj_bullets = extract_bullet_points("\n".join(sections["projects"]))
    
    all_extracted_bullets = exp_bullets + proj_bullets

    # 4. Create Master Candidate Profile JSON schema
    master_profile = {
        "file_source": path.name,
        "raw_character_count": len(raw_text),
        "sections_found": [k for k, v in sections.items() if len(v) > 0],
        "extracted_skills_raw": "\n".join(sections["skills"]),
        "master_bullet_points": [
            {"id": idx + 1, "text": bullet} 
            for idx, bullet in enumerate(all_extracted_bullets)
        ]
    }

    return master_profile

if __name__ == "__main__":
    print("\n--- Running Milestone 3: Ingestor & Parser Test ---\n")
    
    # Create a mock text resume for demonstration test
    sample_resume = """
    ALEX MORGAN
    alex@example.com | 555-0199 | github.com/alex
    
    WORK EXPERIENCE
    - Architected and deployed high-throughput REST APIs using Python, FastAPI, and Pydantic, reducing response latency by 35%.
    - Optimized complex PostgreSQL queries and redesigned database schemas, improving query execution time by 50%.
    - Built asynchronous background task processing workflows using Celery, Redis, and Python.
    - Containerized 12+ backend microservices using Docker and orchestrated deployments on AWS EKS.
    
    TECHNICAL SKILLS
    Python, FastAPI, PostgreSQL, Docker, AWS, Kubernetes, CI/CD, Redis, Git
    
    EDUCATION
    B.S. in Computer Science - University of Technology
    """
    
    # Save temporary mock docx for test run
    test_file = Path("temp_test_resume.docx")
    doc = docx.Document()
    for line in sample_resume.split("\n"):
        doc.add_paragraph(line)
    doc.save(test_file)

    try:
        parsed_data = parse_resume_to_structure(str(test_file))
        print("=== PARSED MASTER PROFILE JSON ===")
        print(f"File Source      : {parsed_data['file_source']}")
        print(f"Sections Found   : {parsed_data['sections_found']}")
        print(f"Bullets Extracted: {len(parsed_data['master_bullet_points'])}\n")
        
        print("Sample Extracted Bullets:")
        for b in parsed_data['master_bullet_points']:
            print(f"  [ID {b['id']}] {b['text']}")
            
    finally:
        # Cleanup temp file
        if test_file.exists():
            test_file.unlink()