import docx
from src.bullet_bank import get_matching_bullets

ACRONYM_MAP = {
    "api": "API", "apis": "APIs", "aws": "AWS", "sql": "SQL",
    "rest": "REST", "ci/cd": "CI/CD", "docker": "Docker", "ui/ux": "UI/UX",
    "json": "JSON", "html": "HTML", "css": "CSS", "gcp": "GCP",
    "ml": "ML", "ai": "AI"
}

def format_keyword(kw: str) -> str:
    words = kw.lower().strip().split()
    formatted_words = []
    for word in words:
        if word in ACRONYM_MAP:
            formatted_words.append(ACRONYM_MAP[word])
        else:
            formatted_words.append(word.title())
    return " ".join(formatted_words)

def inject_keywords_into_resume(docx_path: str, missing_keywords: list[str], output_path: str):
    doc = docx.Document(docx_path)
    
    # 1. Inject missing keywords into the Skills section
    if missing_keywords:
        formatted_keywords = [format_keyword(kw) for kw in missing_keywords]
        keywords_str = ", ".join(formatted_keywords)
        
        target_headers = [
            "skills", "core competencies", "technical skills", 
            "areas of expertise", "technologies", "languages"
        ]
        skills_header_idx = -1
        
        for i, p in enumerate(doc.paragraphs):
            text_lower = p.text.lower().strip()
            if any(header in text_lower for header in target_headers) and len(text_lower.split()) < 5:
                skills_header_idx = int(i)
                break
                
        if skills_header_idx != -1 and skills_header_idx + 1 < len(doc.paragraphs):
            target_para = doc.paragraphs[skills_header_idx + 1]
            if target_para.text.strip():
                target_para.add_run(f", {keywords_str}")
            else:
                target_para.add_run(keywords_str)
        else:
            insert_idx = min(2, len(doc.paragraphs) - 1)
            target_para = doc.paragraphs[insert_idx]
            new_para = target_para.insert_paragraph_before()
            run_bold = new_para.add_run("Core Competencies: ")
            run_bold.bold = True
            new_para.add_run(keywords_str)

    # 2. Select and inject matching professional bullet points into the PROJECTS section
    if missing_keywords:
        matching_bullets = get_matching_bullets(missing_keywords, max_bullets=2)
        project_headers = ["projects", "personal projects", "key projects", "portfolio", "academic projects"]
        project_header_idx = -1
        
        for i, p in enumerate(doc.paragraphs):
            text_lower = p.text.lower().strip()
            if any(h in text_lower for h in project_headers) and len(text_lower.split()) < 5:
                project_header_idx = int(i)
                break
                
        if project_header_idx != -1 and project_header_idx + 1 < len(doc.paragraphs):
            # Found the Projects section: insert bullets right below the header
            target_p = doc.paragraphs[project_header_idx + 1]
            for bullet in matching_bullets:
                target_p.insert_paragraph_before(f"• {bullet}")
        else:
            # Fallback: If no Projects section exists in the resume, create a clean one at the bottom
            new_sec_heading = doc.add_paragraph()
            run_h = new_sec_heading.add_run("PROJECTS")
            run_h.bold = True
            for bullet in matching_bullets:
                doc.add_paragraph(f"• {bullet}")

    doc.save(output_path)
    print("✅ Successfully injected keywords into Skills and achievement bullets into Projects.")