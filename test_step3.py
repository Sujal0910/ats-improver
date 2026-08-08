from src.injector import inject_keywords_into_resume
import os

# Ensure the data directory exists for our outputs
os.makedirs("data", exist_ok=True)

input_resume = "data/sample_resume.docx"
output_resume = "data/optimized_resume.docx"

# Simulated missing keywords from Step 2
missing_from_ats = ["Docker", "REST APIs", "Kubernetes", "Agile Methodologies"]

# If you haven't created a sample doc yet, this script will generate a basic one for testing
import docx
if not os.path.exists(input_resume):
    print("Creating a dummy resume for testing...")
    doc = docx.Document()
    doc.add_paragraph("John Doe\njohn.doe@email.com")
    doc.add_paragraph("Professional Summary\nExperienced engineer.")
    doc.add_paragraph("Technical Skills")
    doc.add_paragraph("Python, AWS, SQL")
    doc.save(input_resume)

print("--- Running ATS Format Injector ---")
inject_keywords_into_resume(input_resume, missing_from_ats, output_resume)

print(f"Done! Open '{output_resume}' in Microsoft Word to see the magic.")